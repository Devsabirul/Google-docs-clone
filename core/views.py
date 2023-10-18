from django.shortcuts import render, redirect
from .models import *
import docx


def home(request):
    files = File_upload.objects.all()
    context = {
        'files': files
    }
    return render(request, 'index.html', context)


def editfile(request, id):
    files = File_upload.objects.get(id=id)
    file_path = files.file_upload.path  # show file path
    showtext = showText(file_path)
    if request.method == "POST":
        # Get existing or updated text
        textupdate = request.POST.get('textupdate')
        writeText(file_path, showtext, textupdate)  # Call Write Text fun
        return redirect('home')
    return render(request, 'editfile.html', {'texts': showtext})


def showText(filepath):
    doc = docx.Document(filepath)  # Open and Read file
    for paragraph in doc.paragraphs:
        text = paragraph.text
    return text


def writeText(filepath, readtext, writetext):
    doc = docx.Document(filepath)  # Open and Read file
    search_text = readtext  # Your existing text
    replacement_text = writetext  # Your updated text
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(
                search_text, replacement_text)

    # Save the modified document
    doc.save(filepath)
    return True
